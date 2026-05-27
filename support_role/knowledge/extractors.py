"""Text extraction for the knowledge base.

Each supported format returns a plain UTF-8 string. Empty / unreadable
files return "". No third-party imports beyond what's listed in
requirements.txt for the matching extension.
"""

from __future__ import annotations

import json
import logging
import re
from dataclasses import dataclass
from pathlib import Path

log = logging.getLogger(__name__)


SUPPORTED_SUFFIXES: tuple[str, ...] = (".pdf", ".docx", ".txt", ".md", ".jsonl")


@dataclass(frozen=True)
class SemanticChunk:
    text: str
    title: str = ""


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    try:
        if suffix == ".pdf":
            return _extract_pdf(path)
        if suffix == ".docx":
            return _extract_docx(path)
        if suffix in (".txt", ".md"):
            return path.read_text(encoding="utf-8", errors="replace")
        if suffix == ".jsonl":
            return _extract_jsonl(path)
    except Exception:
        log.exception("Failed to extract text from %s", path)
        return ""
    log.warning("Unsupported file type: %s", path)
    return ""


def _extract_pdf(path: Path) -> str:
    # pypdf is pure-Python and small. Lazy-import so the rest of the app
    # doesn't pay the import cost when no PDFs are present.
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    parts: list[str] = []
    for page in reader.pages:
        try:
            txt = page.extract_text() or ""
        except Exception:
            txt = ""
        if txt:
            parts.append(txt)
    return "\n\n".join(parts)


def _extract_docx(path: Path) -> str:
    from docx import Document  # python-docx

    doc = Document(str(path))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_jsonl(path: Path) -> str:
    blocks: list[str] = []
    with path.open("r", encoding="utf-8", errors="replace") as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            try:
                obj = json.loads(line)
            except json.JSONDecodeError:
                blocks.append(line)
                continue
            heading = str(obj.get("heading") or obj.get("title") or obj.get("id") or "").strip()
            text = str(obj.get("text") or obj.get("content") or "").strip()
            if not text:
                continue
            if heading:
                blocks.append(f"## {heading}\n\n{text}")
            else:
                blocks.append(text)
    return "\n\n".join(blocks)


def semantic_chunk_text(
    text: str,
    *,
    chunk_chars: int,
    overlap_chars: int,
) -> list[SemanticChunk]:
    """Split text into focused semantic chunks.

    The splitter prefers explicit Markdown headings, then paragraph blocks.
    Character-window splitting is only used inside oversized semantic blocks.
    """
    text = _normalize_text(text)
    if not text:
        return []

    sections = _split_markdown_sections(text)
    chunks: list[SemanticChunk] = []
    for title, section_text in sections:
        section_text = section_text.strip()
        if not section_text:
            continue
        if len(section_text) <= chunk_chars:
            chunks.append(SemanticChunk(text=section_text, title=title))
            continue
        chunks.extend(
            SemanticChunk(text=part, title=title)
            for part in _split_oversized_block(
                section_text,
                chunk_chars=chunk_chars,
                overlap_chars=overlap_chars,
            )
        )
    return deduplicate_chunks(chunks)


def chunk_text(
    text: str,
    chunk_chars: int,
    overlap_chars: int,
) -> list[str]:
    """Backward-compatible wrapper returning only chunk text."""
    return [
        chunk.text
        for chunk in semantic_chunk_text(
            text,
            chunk_chars=chunk_chars,
            overlap_chars=overlap_chars,
        )
    ]


def deduplicate_chunks(chunks: list[SemanticChunk], threshold: float = 0.92) -> list[SemanticChunk]:
    out: list[SemanticChunk] = []
    seen_exact: set[str] = set()
    seen_tokens: list[set[str]] = []
    for chunk in chunks:
        normalized = _fingerprint(chunk.text)
        if not normalized or normalized in seen_exact:
            continue
        tokens = set(normalized.split())
        if _is_near_duplicate(tokens, seen_tokens, threshold):
            continue
        seen_exact.add(normalized)
        seen_tokens.append(tokens)
        out.append(chunk)
    return out


def _normalize_text(text: str) -> str:
    text = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _split_markdown_sections(text: str) -> list[tuple[str, str]]:
    lines = text.splitlines()
    sections: list[tuple[str, list[str]]] = []
    current_title = ""
    current: list[str] = []

    for line in lines:
        heading = re.match(r"^\s{0,3}(#{1,6})\s+(.+?)\s*$", line)
        if heading:
            if current:
                sections.append((current_title, current))
            current_title = heading.group(2).strip()
            current = [line]
            continue
        current.append(line)

    if current:
        sections.append((current_title, current))

    if len(sections) == 1 and not sections[0][0]:
        return _split_paragraph_sections(text)

    return [(title, "\n".join(body).strip()) for title, body in sections]


def _split_paragraph_sections(text: str) -> list[tuple[str, str]]:
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    if not paragraphs:
        return []
    sections: list[tuple[str, str]] = []
    current: list[str] = []
    for paragraph in paragraphs:
        if current and _looks_like_new_topic(paragraph):
            sections.append(("", "\n\n".join(current)))
            current = [paragraph]
        else:
            current.append(paragraph)
    if current:
        sections.append(("", "\n\n".join(current)))
    return sections


def _looks_like_new_topic(paragraph: str) -> bool:
    first_line = paragraph.splitlines()[0].strip()
    if len(first_line) > 90:
        return False
    if first_line.endswith(":"):
        return True
    return bool(re.match(r"^(\d+[\).]|[A-Z][A-Za-z0-9 /&-]{4,80})$", first_line))


def _split_oversized_block(
    text: str,
    *,
    chunk_chars: int,
    overlap_chars: int,
) -> list[str]:
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    if len(paragraphs) <= 1:
        return _split_long_paragraph(text, chunk_chars, overlap_chars)

    chunks: list[str] = []
    current: list[str] = []
    current_len = 0
    for paragraph in paragraphs:
        projected = current_len + len(paragraph) + (2 if current else 0)
        if current and projected > chunk_chars:
            chunks.append("\n\n".join(current).strip())
            current = _paragraph_overlap(current, overlap_chars)
            current_len = sum(len(p) for p in current) + max(0, len(current) - 1) * 2
        if len(paragraph) > chunk_chars:
            if current:
                chunks.append("\n\n".join(current).strip())
                current = []
                current_len = 0
            chunks.extend(_split_long_paragraph(paragraph, chunk_chars, overlap_chars))
            continue
        current.append(paragraph)
        current_len += len(paragraph) + (2 if len(current) > 1 else 0)

    if current:
        chunks.append("\n\n".join(current).strip())
    return [c for c in chunks if c]


def _paragraph_overlap(paragraphs: list[str], overlap_chars: int) -> list[str]:
    if overlap_chars <= 0 or not paragraphs:
        return []
    out: list[str] = []
    used = 0
    for paragraph in reversed(paragraphs):
        if used + len(paragraph) > overlap_chars:
            break
        out.append(paragraph)
        used += len(paragraph)
    return list(reversed(out[-1:]))


def _split_long_paragraph(text: str, chunk_chars: int, overlap_chars: int) -> list[str]:
    chunks: list[str] = []
    start = 0
    n = len(text)
    while start < n:
        end = min(start + chunk_chars, n)
        if end < n:
            for sep in (". ", "? ", "! ", "; ", ", "):
                idx = text.rfind(sep, start + chunk_chars // 2, end)
                if idx != -1:
                    end = idx + len(sep)
                    break
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= n:
            break
        start = max(0, end - min(overlap_chars, chunk_chars // 5))
    return chunks


def _fingerprint(text: str) -> str:
    lowered = text.lower()
    lowered = re.sub(r"```.*?```", " ", lowered, flags=re.DOTALL)
    lowered = re.sub(r"[^a-z0-9]+", " ", lowered)
    return " ".join(lowered.split())


def _is_near_duplicate(tokens: set[str], seen_tokens: list[set[str]], threshold: float) -> bool:
    if len(tokens) < 8:
        return False
    for existing in seen_tokens:
        smaller = min(len(tokens), len(existing))
        larger = max(len(tokens), len(existing))
        if smaller == 0 or smaller / larger < 0.65:
            continue
        overlap = len(tokens & existing) / len(tokens | existing)
        if overlap >= threshold:
            return True
    return False
